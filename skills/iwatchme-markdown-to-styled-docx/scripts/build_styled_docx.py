#!/usr/bin/env python3
"""
build_styled_docx.py — Markdown 简历 → 样式化 .docx

使用 YAML frontmatter 作为唯一的头部联系信息来源。

Usage:
    pip install python-docx --break-system-packages
    python3 build_styled_docx.py input.md output.docx
"""

import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE = RGBColor(0x17, 0x59, 0xA0)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x26, 0x26, 0x26)

FONT_LATIN = "Arial"
FONT_CJK = "微软雅黑"


def set_run_font(run, size=10.5, bold=False, color=BLACK, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = FONT_LATIN
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), FONT_LATIN)
    r_fonts.set(qn("w:hAnsi"), FONT_LATIN)
    r_fonts.set(qn("w:eastAsia"), FONT_CJK)
    r_fonts.set(qn("w:cs"), FONT_CJK)


def add_run(paragraph, text, bold=False, size=10.5, color=BLACK, italic=False):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    return run


def set_paragraph_spacing(paragraph, before=0, after=0, line=None):
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line is not None:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")


def set_indent(paragraph, left_twips=0, hanging_twips=0):
    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.set(qn("w:left"), str(left_twips))
    if hanging_twips:
        ind.set(qn("w:hanging"), str(hanging_twips))


def set_paragraph_pagination(
    paragraph, keep_next=False, keep_lines=False, page_break_before=False
):
    p_pr = paragraph._p.get_or_add_pPr()
    for tag, enabled in (
        ("w:keepNext", keep_next),
        ("w:keepLines", keep_lines),
        ("w:pageBreakBefore", page_break_before),
    ):
        existing = p_pr.find(qn(tag))
        if enabled:
            if existing is None:
                p_pr.append(OxmlElement(tag))
        elif existing is not None:
            p_pr.remove(existing)


def ensure_bullet_numbering(doc):
    numbering_part = doc.part.numbering_part
    num_id = getattr(numbering_part, "_resume_bullet_num_id", None)
    if num_id is not None:
        return num_id

    numbering = numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_id))

    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "multilevel")
    abstract_num.append(multi_level)

    for level, bullet_char, left, hanging, size, latin_font, cjk_font in (
        (0, "●", 360, 360, 20, FONT_LATIN, FONT_CJK),
        (1, "○", 760, 400, 20, FONT_LATIN, FONT_CJK),
    ):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))

        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)

        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet")
        lvl.append(num_fmt)

        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), bullet_char)
        lvl.append(lvl_text)

        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)

        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)

        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.append(ind)
        lvl.append(p_pr)

        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), latin_font)
        r_fonts.set(qn("w:hAnsi"), latin_font)
        r_fonts.set(qn("w:eastAsia"), cjk_font)
        r_fonts.set(qn("w:cs"), cjk_font)
        r_pr.append(r_fonts)

        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(size))
        r_pr.append(sz)

        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), str(size))
        r_pr.append(sz_cs)
        lvl.append(r_pr)

        abstract_num.append(lvl)

    numbering.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)

    numbering_part._resume_bullet_num_id = num_id
    return num_id


def apply_bullet_level(doc, paragraph, level, left_twips, hanging_twips):
    num_id = ensure_bullet_numbering(doc)
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)

    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), str(level))

    num_id_el = num_pr.find(qn("w:numId"))
    if num_id_el is None:
        num_id_el = OxmlElement("w:numId")
        num_pr.append(num_id_el)
    num_id_el.set(qn("w:val"), str(num_id))

    set_indent(paragraph, left_twips=left_twips, hanging_twips=hanging_twips)


def add_paragraph_bottom_border(paragraph, color="1759A0", sz="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    existing = p_pr.find(qn("w:pBdr"))
    if existing is not None:
        p_pr.remove(existing)
    p_pr.append(p_bdr)


def add_tab_stop(paragraph, position_cm, alignment="right"):
    p_pr = paragraph._p.get_or_add_pPr()
    tabs = p_pr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        p_pr.append(tabs)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), alignment)
    tab.set(qn("w:pos"), str(int(position_cm * 567)))
    tabs.append(tab)


def set_cell_borders_none(cell):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def parse_frontmatter(content):
    if not content.startswith("---"):
        return {}, content

    end_idx = content.find("\n---", 3)
    if end_idx == -1:
        return {}, content

    fm_block = content[3:end_idx].strip()
    body = content[end_idx + 4 :].lstrip("\n")

    metadata = {}
    for line in fm_block.splitlines():
        if ":" in line:
            key, _, value = line.partition(":")
            key = key.strip()
            value = value.strip()
            if key and value:
                metadata[key] = value

    return metadata, body


def parse_job_title(text):
    match = re.match(r"^(.+?)\s*[·•]\s*(.+?)（(.+?)）\s*$", text)
    if match:
        return match.group(1).strip(), match.group(2).strip(), match.group(3).strip()
    return text, "", ""


def parse_structured_row(text, line_no):
    parts = [part.strip() for part in text.split("|")]
    if len(parts) not in (2, 3) or not all(parts):
        raise ValueError(
            f"Line {line_no}: structured row must be 'left | right' or 'left | middle | right', got: {text}"
        )
    return parts


def parse_line_indent(line, line_no):
    expanded = line.expandtabs(2)
    stripped = expanded.lstrip(" ")
    indent = len(expanded) - len(stripped)
    if indent % 2 != 0:
        raise ValueError(f"Line {line_no}: indentation must use multiples of 2 spaces")
    return indent // 2, stripped


def parse_markdown(content):
    tokens = []
    entry_active = False
    entry_has_subtitle = False
    entry_has_children = False
    current_list_level = None
    current_entry_token = None

    def close_entry():
        nonlocal entry_active, entry_has_subtitle, entry_has_children
        nonlocal current_list_level, current_entry_token
        entry_active = False
        entry_has_subtitle = False
        entry_has_children = False
        current_list_level = None
        current_entry_token = None

    for line_no, line in enumerate(content.split("\n"), 1):
        stripped = line.strip()
        if not stripped:
            continue

        indent_level, stripped = parse_line_indent(line, line_no)

        if stripped.startswith("# ") and not stripped.startswith("## "):
            tokens.append({"type": "name", "text": stripped[2:].strip()})
            close_entry()
        elif stripped.startswith("## "):
            tokens.append({"type": "section", "text": stripped[3:].strip()})
            close_entry()
        elif stripped == "---":
            continue
        elif stripped.startswith("- "):
            text = stripped[2:].strip()
            if indent_level == 0:
                close_entry()
                if "|" in text:
                    parts = parse_structured_row(text, line_no)
                    tokens.append(
                        {
                            "type": "structured_row",
                            "parts": parts,
                            "has_details": False,
                        }
                    )
                    current_entry_token = tokens[-1]
                    entry_active = True
                elif parse_job_title(text)[1]:
                    title, company, date = parse_job_title(text)
                    tokens.append(
                        {
                            "type": "structured_row",
                            "parts": [title, company, date],
                            "has_details": False,
                        }
                    )
                    current_entry_token = tokens[-1]
                    entry_active = True
                else:
                    tokens.append(
                        {
                            "type": "list_item",
                            "text": text,
                            "level": 0,
                        }
                    )
                    current_list_level = 0
            elif indent_level == 1:
                if not entry_active:
                    raise ValueError(
                        f"Line {line_no}: nested list item must belong to a structured row"
                    )
                tokens.append({"type": "nested_bullet", "text": text, "level": 1})
                entry_has_children = True
                if current_entry_token is not None:
                    current_entry_token["has_details"] = True
                current_list_level = 1
            elif indent_level == 2:
                if not entry_active or current_list_level not in (1, 2):
                    raise ValueError(
                        f"Line {line_no}: level-3 detail bullets must belong to a level-2 list item"
                    )
                tokens.append({"type": "nested_bullet", "text": text, "level": 2})
                current_list_level = 2
            else:
                raise ValueError(
                    f"Line {line_no}: structured sections support at most 3 list levels"
                )
        elif indent_level >= 1:
            if not entry_active:
                raise ValueError(
                    f"Line {line_no}: indented text must belong to a structured row"
                )
            if entry_has_children or entry_has_subtitle:
                raise ValueError(
                    f"Line {line_no}: only one subtitle line is allowed before nested bullets"
                )
            tokens.append({"type": "subtitle", "text": stripped})
            entry_has_subtitle = True
            if current_entry_token is not None:
                current_entry_token["has_details"] = True
        elif stripped.startswith("### "):
            tokens.append({"type": "job_title", "text": stripped[4:].strip()})
        elif stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            inner = stripped[2:-2].strip()
            tokens.append({"type": "subtitle", "text": inner})
        elif stripped.startswith("- "):
            tokens.append(
                {
                    "type": "list_item",
                    "text": stripped[2:].strip(),
                    "level": 0,
                }
            )
        else:
            tokens.append({"type": "paragraph", "text": stripped})
    return tokens


def build_header(doc, name, email, phone, city):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table.rows[0].cells:
        set_cell_borders_none(cell)

    left_cell = table.cell(0, 0)
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    left_paragraph = left_cell.paragraphs[0]
    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(left_paragraph, 0, 0)
    email_run = add_run(left_paragraph, email, size=9, color=BLUE)
    email_run.font.underline = True
    if city:
        left_paragraph.add_run("\n")
        add_run(left_paragraph, city, size=9)

    middle_cell = table.cell(0, 1)
    middle_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    middle_paragraph = middle_cell.paragraphs[0]
    middle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(middle_paragraph, 0, 0)
    add_run(middle_paragraph, name, bold=True, size=20, color=DARK_GRAY)

    right_cell = table.cell(0, 2)
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    right_paragraph = right_cell.paragraphs[0]
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(right_paragraph, 0, 0)
    add_run(right_paragraph, phone, size=9)

    for cell in table.rows[0].cells:
        cell.width = Cm(5.5)


def add_section_heading(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=140, after=40)
    set_paragraph_pagination(paragraph, keep_next=True, keep_lines=False)
    add_run(paragraph, text, bold=True, size=13, color=BLUE)
    add_paragraph_bottom_border(paragraph)


def add_structured_header(doc, parts, right_align=True):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=100, after=20)
    add_run(paragraph, parts[0], bold=True, size=11.5)
    if len(parts) == 2:
        add_tab_stop(paragraph, 16.2, "right" if right_align else "left")
        paragraph.add_run("\t")
        add_run(paragraph, parts[1], bold=False, size=11.5)
    elif len(parts) == 3:
        add_tab_stop(paragraph, 8.5, "center" if right_align else "left")
        add_tab_stop(paragraph, 16.2, "right" if right_align else "left")
        paragraph.add_run("\t")
        add_run(paragraph, parts[1], bold=True, size=11.5)
        paragraph.add_run("\t")
        add_run(paragraph, parts[2], bold=True, size=11.5)


def add_subtitle_line(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=0, after=50)
    add_run(paragraph, text, size=9.5, color=DARK_GRAY)


def add_labeled_bullet(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=50, after=20)
    apply_bullet_level(doc, paragraph, level=0, left_twips=360, hanging_twips=360)
    match = re.match(r"^([^：]+：)(.*)", text)
    if match:
        add_run(paragraph, match.group(1), bold=True, size=10.5)
        add_run(paragraph, match.group(2), bold=False, size=10.5)
    else:
        add_inline_bold(paragraph, text, size=10.5)


def add_detail_bullet(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=16, after=16)
    apply_bullet_level(doc, paragraph, level=1, left_twips=760, hanging_twips=400)
    add_inline_bold(paragraph, text, size=10)


def add_simple_bullet(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=20, after=20)
    apply_bullet_level(doc, paragraph, level=0, left_twips=360, hanging_twips=360)
    add_inline_bold(paragraph, text, size=10.5)


def add_plain_line(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=20, after=20)
    add_inline_bold(paragraph, text, size=10.5)


def add_compact_structured_item(doc, parts):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=20, after=20)
    apply_bullet_level(doc, paragraph, level=0, left_twips=360, hanging_twips=360)
    add_run(paragraph, f"{parts[0]}：", bold=True, size=10.5)
    add_run(paragraph, parts[1], bold=False, size=10.5)


def add_summary(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=80, after=80)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(paragraph, text, size=10.5)


def add_inline_bold(paragraph, text, size=10.5):
    parts = re.split(r"\*\*(.*?)\*\*", text)
    for idx, part in enumerate(parts):
        if part:
            add_run(paragraph, part, bold=(idx % 2 == 1), size=size)


def build_docx(md_content, output_path):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.9)
        section.right_margin = Cm(1.9)

    style = doc.styles["Normal"]
    style.font.name = FONT_LATIN
    style.font.size = Pt(10.5)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), FONT_CJK)

    metadata, body = parse_frontmatter(md_content)
    fm_name = metadata.get("name", "")
    fm_email = metadata.get("email", "")
    fm_phone = metadata.get("phone", "")
    fm_city = metadata.get("city", "")

    if not fm_name or not fm_email:
        raise ValueError("frontmatter must include at least 'name' and 'email'")

    tokens = parse_markdown(body)

    build_header(doc, fm_name, fm_email, fm_phone, fm_city)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=60)

    nested_context_level = None

    for token in tokens:
        token_type = token["type"]
        text = token.get("text", "")

        if token_type == "name":
            continue
        elif token_type == "paragraph":
            add_summary(doc, text)
        elif token_type == "section":
            add_section_heading(doc, text)
            nested_context_level = None
        elif token_type == "job_title":
            title, company, date = parse_job_title(text)
            add_structured_header(doc, [title, company, date])
            nested_context_level = None
        elif token_type == "structured_row":
            if len(token["parts"]) == 2 and not token.get("has_details", False):
                add_compact_structured_item(doc, token["parts"])
            else:
                add_structured_header(doc, token["parts"], right_align=True)
            nested_context_level = None
        elif token_type == "subtitle":
            add_subtitle_line(doc, text)
        elif token_type == "nested_bullet":
            if token["level"] == 1:
                add_labeled_bullet(doc, text)
            else:
                add_detail_bullet(doc, text)
            nested_context_level = token["level"]
        elif token_type == "list_item":
            if nested_context_level == 1:
                add_detail_bullet(doc, text)
            else:
                add_simple_bullet(doc, text)

    doc.save(output_path)
    print(f"[OK] 已生成: {output_path}")


def main():
    if len(sys.argv) < 3:
        print("Usage: python3 build_styled_docx.py input.md output.docx")
        sys.exit(1)
    with open(sys.argv[1], "r", encoding="utf-8") as handle:
        content = handle.read()
    build_docx(content, sys.argv[2])


if __name__ == "__main__":
    main()
