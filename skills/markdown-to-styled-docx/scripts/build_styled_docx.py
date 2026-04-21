#!/usr/bin/env python3
"""
build_styled_docx.py — Markdown 简历 → 样式化 .docx

支持 YAML frontmatter（推荐）和 📧/📱 行（兼容旧格式）两种联系信息写法。

Usage:
    pip install python-docx --break-system-packages
    python3 build_styled_docx.py input.md output.docx
"""

import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE = RGBColor(0x17, 0x59, 0xA0)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x26, 0x26, 0x26)

FONT_LATIN = "Arial"
FONT_CJK = "微软雅黑"


def set_run_font(run, size=10.5, bold=False, color=BLACK, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = FONT_LATIN
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), FONT_LATIN)
    r_fonts.set(qn("w:hAnsi"), FONT_LATIN)
    r_fonts.set(qn("w:eastAsia"), FONT_CJK)
    r_fonts.set(qn("w:cs"), FONT_CJK)


def add_run(paragraph, text, bold=False, size=10.5, color=BLACK, italic=False):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    return run


def set_paragraph_spacing(paragraph, before=0, after=0, line=None):
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line is not None:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")


def set_indent(paragraph, left_twips=0, hanging_twips=0):
    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.set(qn("w:left"), str(left_twips))
    if hanging_twips:
        ind.set(qn("w:hanging"), str(hanging_twips))


def add_paragraph_bottom_border(paragraph, color="1759A0", sz="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    existing = p_pr.find(qn("w:pBdr"))
    if existing is not None:
        p_pr.remove(existing)
    p_pr.append(p_bdr)


def add_tab_stop(paragraph, position_cm, alignment="right"):
    p_pr = paragraph._p.get_or_add_pPr()
    tabs = p_pr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        p_pr.append(tabs)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), alignment)
    tab.set(qn("w:pos"), str(int(position_cm * 567)))
    tabs.append(tab)


def set_cell_borders_none(cell):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def parse_frontmatter(content):
    if not content.startswith("---"):
        return {}, content

    end_idx = content.find("\n---", 3)
    if end_idx == -1:
        return {}, content

    fm_block = content[3:end_idx].strip()
    body = content[end_idx + 4 :].lstrip("\n")

    metadata = {}
    for line in fm_block.splitlines():
        if ":" in line:
            key, _, value = line.partition(":")
            key = key.strip()
            value = value.strip()
            if key and value:
                metadata[key] = value

    return metadata, body


def parse_contact(text):
    parts = [part.strip() for part in re.split(r"\|", text)]
    email = phone = city = ""
    for part in parts:
        part = re.sub(r"[📧📱]", "", part).strip()
        if "@" in part:
            email = part
        elif re.match(r"[\+\d]", part):
            phone = part
        elif "意向城市" in part:
            city = part.replace("意向城市：", "").strip()
    return email, phone, city


def parse_job_title(text):
    match = re.match(r"^(.+?)\s*[·•]\s*(.+?)（(.+?)）\s*$", text)
    if match:
        return match.group(1).strip(), match.group(2).strip(), match.group(3).strip()
    return text, "", ""


def parse_job_row(text, line_no):
    parts = [part.strip() for part in text.split("|")]
    if len(parts) != 3 or not all(parts):
        raise ValueError(
            f"Line {line_no}: job row must be 'title | company | date', got: {text}"
        )
    return parts[0], parts[1], parts[2]


def is_work_experience_section(section_name):
    return "工作经历" in section_name


def parse_line_indent(line, line_no):
    expanded = line.expandtabs(2)
    stripped = expanded.lstrip(" ")
    indent = len(expanded) - len(stripped)
    if indent % 2 != 0:
        raise ValueError(f"Line {line_no}: indentation must use multiples of 2 spaces")
    return indent // 2, stripped


def parse_markdown(content):
    tokens = []
    current_section = ""
    in_work_section = False
    job_active = False
    job_has_subtitle = False
    job_has_children = False
    current_child_kind = None

    for line_no, line in enumerate(content.split("\n"), 1):
        stripped = line.strip()
        if not stripped:
            continue

        indent_level, stripped = parse_line_indent(line, line_no)

        if stripped.startswith("# ") and not stripped.startswith("## "):
            tokens.append({"type": "name", "text": stripped[2:].strip()})
            job_active = False
            current_child_kind = None
        elif stripped.startswith("## "):
            current_section = stripped[3:].strip()
            in_work_section = is_work_experience_section(current_section)
            tokens.append({"type": "section", "text": current_section})
            job_active = False
            job_has_subtitle = False
            job_has_children = False
            current_child_kind = None
        elif stripped == "---":
            continue
        elif re.match(r"^[📧📱]", stripped) or ("|" in stripped and "@" in stripped):
            tokens.append({"type": "contact", "text": stripped})
        elif in_work_section and stripped.startswith("- "):
            text = stripped[2:].strip()
            if indent_level == 0:
                if "|" in text:
                    title, company, date = parse_job_row(text, line_no)
                    tokens.append(
                        {
                            "type": "job_row",
                            "title": title,
                            "company": company,
                            "date": date,
                        }
                    )
                else:
                    # Backward compatibility for existing notes while the new
                    # list-driven syntax becomes the documented path.
                    title, company, date = parse_job_title(text)
                    tokens.append(
                        {
                            "type": "job_row",
                            "title": title,
                            "company": company,
                            "date": date,
                        }
                    )
                job_active = True
                job_has_subtitle = False
                job_has_children = False
                current_child_kind = None
            elif indent_level == 1:
                if not job_active:
                    raise ValueError(
                        f"Line {line_no}: nested list item must belong to a job row"
                    )
                if re.match(r"^(核心项目|基础项目)：", text):
                    tokens.append({"type": "project_bullet", "text": text})
                    current_child_kind = "project"
                else:
                    tokens.append({"type": "simple_bullet", "text": text})
                    current_child_kind = "simple"
                job_has_children = True
            elif indent_level == 2:
                if current_child_kind != "project":
                    raise ValueError(
                        f"Line {line_no}: level-3 detail bullets must belong to a core/basic project"
                    )
                tokens.append({"type": "detail_bullet", "text": text})
            else:
                raise ValueError(
                    f"Line {line_no}: work experience supports at most 3 list levels"
                )
        elif in_work_section and indent_level >= 1:
            if not job_active:
                raise ValueError(
                    f"Line {line_no}: indented text in work experience must belong to a job row"
                )
            if job_has_children or job_has_subtitle:
                raise ValueError(
                    f"Line {line_no}: only one subtitle line is allowed before nested bullets"
                )
            tokens.append({"type": "subtitle", "text": stripped})
            job_has_subtitle = True
        elif stripped.startswith("### "):
            tokens.append({"type": "job_title", "text": stripped[4:].strip()})
        elif stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            inner = stripped[2:-2].strip()
            if re.match(r"^(核心项目|基础项目)：", inner):
                tokens.append({"type": "project_bullet", "text": inner})
            else:
                tokens.append({"type": "subtitle", "text": inner})
        elif stripped.startswith("- "):
            tokens.append({"type": "list_item", "text": stripped[2:].strip()})
        else:
            tokens.append({"type": "paragraph", "text": stripped})
    return tokens


def build_header(doc, name, email, phone, city):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table.rows[0].cells:
        set_cell_borders_none(cell)

    left_cell = table.cell(0, 0)
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    left_paragraph = left_cell.paragraphs[0]
    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(left_paragraph, 0, 0)
    email_run = add_run(left_paragraph, email, size=9, color=BLUE)
    email_run.font.underline = True
    if city:
        left_paragraph.add_run("\n")
        add_run(left_paragraph, f"意向城市：{city}", size=9)

    middle_cell = table.cell(0, 1)
    middle_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    middle_paragraph = middle_cell.paragraphs[0]
    middle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(middle_paragraph, 0, 0)
    add_run(middle_paragraph, name, bold=True, size=20, color=DARK_GRAY)

    right_cell = table.cell(0, 2)
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    right_paragraph = right_cell.paragraphs[0]
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(right_paragraph, 0, 0)
    add_run(right_paragraph, phone, size=9)

    for cell in table.rows[0].cells:
        cell.width = Cm(5.5)


def add_section_heading(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=140, after=40)
    add_run(paragraph, text, bold=True, size=13, color=BLUE)
    add_paragraph_bottom_border(paragraph)


def add_job_header(doc, title, company, date):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=100, after=20)
    add_tab_stop(paragraph, 8.5, "center")
    add_tab_stop(paragraph, 16.2, "right")
    add_run(paragraph, title, bold=True, size=11.5)
    paragraph.add_run("\t")
    add_run(paragraph, company, bold=True, size=11.5)
    paragraph.add_run("\t")
    add_run(paragraph, date, bold=True, size=11.5)


def add_subtitle_line(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=0, after=50)
    add_run(paragraph, text, size=9.5, color=DARK_GRAY)


def add_project_bullet(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=50, after=20)
    set_indent(paragraph, left_twips=360, hanging_twips=360)
    add_run(paragraph, "● ", bold=True, size=10.5)
    match = re.match(r"^((?:核心项目|基础项目)：)(.*)", text)
    if match:
        add_run(paragraph, match.group(1), bold=True, size=10.5)
        add_run(paragraph, match.group(2), bold=False, size=10.5)
    else:
        add_run(paragraph, text, bold=True, size=10.5)


def add_detail_bullet(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=16, after=16)
    set_indent(paragraph, left_twips=720, hanging_twips=360)
    add_run(paragraph, "○ ", size=10)
    add_inline_bold(paragraph, text, size=10)


def add_simple_bullet(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=20, after=20)
    set_indent(paragraph, left_twips=360, hanging_twips=360)
    add_run(paragraph, "● ", size=10.5)
    add_inline_bold(paragraph, text, size=10.5)


def add_summary(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=80, after=80)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(paragraph, text, size=10.5)


def add_inline_bold(paragraph, text, size=10.5):
    parts = re.split(r"\*\*(.*?)\*\*", text)
    for idx, part in enumerate(parts):
        if part:
            add_run(paragraph, part, bold=(idx % 2 == 1), size=size)


def build_docx(md_content, output_path):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.9)
        section.right_margin = Cm(1.9)

    style = doc.styles["Normal"]
    style.font.name = FONT_LATIN
    style.font.size = Pt(10.5)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), FONT_CJK)

    metadata, body = parse_frontmatter(md_content)
    has_frontmatter = bool(metadata)

    fm_name = metadata.get("name", "")
    fm_email = metadata.get("email", "")
    fm_phone = metadata.get("phone", "")
    fm_city = metadata.get("city", "")

    tokens = parse_markdown(body if has_frontmatter else md_content)

    header_built = False
    if has_frontmatter and fm_name and fm_email:
        build_header(doc, fm_name, fm_email, fm_phone, fm_city)
        spacer = doc.add_paragraph()
        set_paragraph_spacing(spacer, before=0, after=60)
        header_built = True

    in_project_context = False
    name = fm_name

    for token in tokens:
        token_type = token["type"]
        text = token.get("text", "")

        if token_type == "name":
            if not header_built:
                name = text
        elif token_type == "contact":
            if not header_built:
                email, phone, city = parse_contact(text)
                if name:
                    build_header(doc, name, email, phone, city)
                    spacer = doc.add_paragraph()
                    set_paragraph_spacing(spacer, before=0, after=60)
                    header_built = True
        elif token_type == "paragraph":
            add_summary(doc, text)
        elif token_type == "section":
            add_section_heading(doc, text)
            in_project_context = False
        elif token_type == "job_title":
            title, company, date = parse_job_title(text)
            add_job_header(doc, title, company, date)
            in_project_context = False
        elif token_type == "job_row":
            add_job_header(doc, token["title"], token["company"], token["date"])
            in_project_context = False
        elif token_type == "subtitle":
            add_subtitle_line(doc, text)
        elif token_type == "project_bullet":
            add_project_bullet(doc, text)
            in_project_context = True
        elif token_type == "simple_bullet":
            add_simple_bullet(doc, text)
            in_project_context = False
        elif token_type == "detail_bullet":
            add_detail_bullet(doc, text)
            in_project_context = True
        elif token_type == "list_item":
            if in_project_context:
                add_detail_bullet(doc, text)
            else:
                add_simple_bullet(doc, text)

    doc.save(output_path)
    print(f"[OK] 已生成: {output_path}")


def main():
    if len(sys.argv) < 3:
        print("Usage: python3 build_styled_docx.py input.md output.docx")
        sys.exit(1)
    with open(sys.argv[1], "r", encoding="utf-8") as handle:
        content = handle.read()
    build_docx(content, sys.argv[2])


if __name__ == "__main__":
    main()
